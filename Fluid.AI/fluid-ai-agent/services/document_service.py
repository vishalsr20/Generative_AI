import os
import re

from docx import Document
from docx.shared import Pt


def create_word_document(content: str, filename="generated_document.docx"):

    os.makedirs("outputs", exist_ok=True)

    doc = Document()

    # Default Font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = content.split("\n")

    for line in lines:

        line = line.strip()

        if not line:
            continue

        # Heading 1
        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)

        # Heading 2
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)

        # Heading 3
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)

        # Bullet
        elif line.startswith("* "):
            doc.add_paragraph(
                line.replace("* ", ""),
                style="List Bullet"
            )

        elif line.startswith("- "):
            doc.add_paragraph(
                line.replace("- ", ""),
                style="List Bullet"
            )

        else:
            doc.add_paragraph(line)

    filepath = os.path.join("outputs", filename)

    doc.save(filepath)

    return filepath